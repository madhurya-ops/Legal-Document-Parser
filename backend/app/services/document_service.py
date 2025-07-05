"""Document Service for LegalDoc application."""

import logging
import hashlib
import uuid
import os
from typing import List, Dict, Any, Optional
from datetime import datetime
from fastapi import UploadFile
from sqlalchemy.orm import Session

from ..exceptions import FileProcessingError, NotFoundError, DuplicateError, ValidationError
from ..models import Document, User
from ..schemas import DocumentCreate, DocumentResponse

logger = logging.getLogger(__name__)


class ExtractedText:
    """Wrapper for extracted text with metadata."""
    
    def __init__(
        self,
        content: str,
        metadata: Optional[Dict[str, Any]] = None,
        page_count: Optional[int] = None,
        word_count: Optional[int] = None
    ):
        self.content = content
        self.metadata = metadata or {}
        self.page_count = page_count
        self.word_count = word_count or len(content.split())


class DocumentService:
    """Document service for file management and processing."""
    
    def __init__(self):
        self.storage_path = "uploads"  # Configure as needed
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.allowed_types = {
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',
            'text/plain'
        }
        self._ensure_storage_path()
    
    def _ensure_storage_path(self):
        """Ensure storage directory exists."""
        if not os.path.exists(self.storage_path):
            os.makedirs(self.storage_path)
    
    async def upload_document(
        self,
        file: UploadFile,
        user_id: str,
        db: Session = None
    ) -> Dict[str, Any]:
        """
        Upload and process document.
        
        Features:
        - File validation
        - Virus scanning (placeholder)
        - Text extraction
        - Metadata extraction
        - Vector embedding (placeholder)
        - Storage management
        """
        try:
            # Validate file
            await self._validate_file(file)
            
            # Calculate file hash for duplicate detection
            file_content = await file.read()
            file_hash = hashlib.sha256(file_content).hexdigest()
            
            # Check for duplicates
            if db:
                existing = db.query(Document).filter(
                    Document.file_hash == file_hash,
                    Document.user_id == uuid.UUID(user_id)
                ).first()
                
                if existing:
                    return {
                        "message": "Document already exists",
                        "document": self._document_to_dict(existing),
                        "is_duplicate": True
                    }
            
            # Generate unique filename
            file_extension = self._get_file_extension(file.filename)
            unique_filename = f"{uuid.uuid4()}{file_extension}"
            file_path = os.path.join(self.storage_path, unique_filename)
            
            # Save file to storage
            with open(file_path, "wb") as f:
                f.write(file_content)
            
            # Extract text and metadata
            extracted_text = await self.extract_text(file, file_content)
            
            # Create document record
            document_data = {
                "filename": unique_filename,
                "original_filename": file.filename,
                "file_hash": file_hash,
                "file_size": str(len(file_content)),
                "file_type": file.content_type,
                "user_id": uuid.UUID(user_id)
            }
            
            if db:
                document = Document(**document_data)
                db.add(document)
                db.commit()
                db.refresh(document)
                
                # Generate embeddings in background (placeholder)
                # await self.generate_embeddings(extracted_text.content)
                
                logger.info(f"Document {document.id} uploaded successfully for user {user_id}")
                
                return {
                    "message": "Document uploaded successfully",
                    "document": self._document_to_dict(document),
                    "is_duplicate": False,
                    "text_content": extracted_text.content[:500] + "..." if len(extracted_text.content) > 500 else extracted_text.content,
                    "metadata": extracted_text.metadata
                }
            else:
                # Return document data without DB storage
                return {
                    "message": "Document processed successfully",
                    "document": document_data,
                    "is_duplicate": False,
                    "text_content": extracted_text.content[:500] + "..." if len(extracted_text.content) > 500 else extracted_text.content,
                    "metadata": extracted_text.metadata
                }
            
        except Exception as e:
            logger.error(f"Error uploading document: {str(e)}")
            raise FileProcessingError(f"Failed to upload document: {str(e)}")
    
    async def extract_text(
        self,
        file: UploadFile,
        file_content: Optional[bytes] = None
    ) -> ExtractedText:
        """
        Extract text from document.
        
        Features:
        - Multi-format support
        - OCR capabilities (placeholder)
        - Structure preservation
        - Metadata extraction
        """
        try:
            if not file_content:
                file_content = await file.read()
            
            content_type = file.content_type
            
            if content_type == 'application/pdf':
                return await self._extract_from_pdf(file_content)
            elif content_type in [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/msword'
            ]:
                return await self._extract_from_docx(file_content)
            elif content_type == 'text/plain':
                return await self._extract_from_text(file_content)
            else:
                raise FileProcessingError(f"Unsupported file type: {content_type}")
            
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            raise FileProcessingError(f"Failed to extract text: {str(e)}")
    
    async def get_document_context(
        self,
        document_ids: List[str],
        query: str,
        max_context_length: int = 2000
    ) -> str:
        """
        Get relevant document context for a query.
        
        Features:
        - Multi-document context
        - Relevance scoring
        - Context optimization
        """
        try:
            # Placeholder implementation
            # In a real implementation, this would:
            # 1. Retrieve documents from database
            # 2. Search through document content using vector similarity
            # 3. Extract most relevant passages
            # 4. Combine and optimize context
            
            context_parts = []
            for doc_id in document_ids[:3]:  # Limit to 3 documents
                # Simulate document retrieval and context extraction
                context_parts.append(f"Context from document {doc_id}: Relevant legal content related to '{query}'.")
            
            full_context = " ".join(context_parts)
            
            # Truncate if too long
            if len(full_context) > max_context_length:
                full_context = full_context[:max_context_length] + "..."
            
            return full_context
            
        except Exception as e:
            logger.error(f"Error getting document context: {str(e)}")
            return ""
    
    async def generate_embeddings(
        self,
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ) -> List[List[float]]:
        """
        Generate vector embeddings.
        
        Features:
        - Chunking strategy
        - Embedding generation
        - Vector storage
        - Similarity search
        """
        try:
            # Placeholder implementation
            # In a real implementation, this would:
            # 1. Split text into chunks with overlap
            # 2. Generate embeddings using OpenAI or similar service
            # 3. Store embeddings in vector database (FAISS/Pinecone)
            # 4. Return embedding vectors
            
            chunks = self._chunk_text(text, chunk_size, chunk_overlap)
            embeddings = []
            
            for chunk in chunks:
                # Simulate embedding generation
                # embedding = await self.embedding_service.embed(chunk)
                embedding = [0.1] * 1536  # Placeholder 1536-dimensional embedding
                embeddings.append(embedding)
            
            logger.info(f"Generated {len(embeddings)} embeddings for text")
            return embeddings
            
        except Exception as e:
            logger.error(f"Error generating embeddings: {str(e)}")
            return []
    
    async def get_user_documents(
        self,
        user_id: str,
        skip: int = 0,
        limit: int = 50,
        search: Optional[str] = None,
        db: Session = None
    ) -> List[Dict[str, Any]]:
        """Get user's documents with optional search."""
        try:
            if not db:
                return []
            
            query = db.query(Document).filter(Document.user_id == uuid.UUID(user_id))
            
            if search:
                query = query.filter(
                    Document.original_filename.ilike(f"%{search}%")
                )
            
            documents = query.offset(skip).limit(limit).all()
            
            return [self._document_to_dict(doc) for doc in documents]
            
        except Exception as e:
            logger.error(f"Error getting user documents: {str(e)}")
            return []
    
    async def get_document(
        self,
        document_id: str,
        user_id: str,
        db: Session = None
    ) -> Optional[Dict[str, Any]]:
        """Get specific document by ID."""
        try:
            if not db:
                return None
            
            document = db.query(Document).filter(
                Document.id == uuid.UUID(document_id),
                Document.user_id == uuid.UUID(user_id)
            ).first()
            
            if not document:
                raise NotFoundError("Document", document_id)
            
            return self._document_to_dict(document)
            
        except Exception as e:
            logger.error(f"Error getting document: {str(e)}")
            return None
    
    async def delete_document(
        self,
        document_id: str,
        user_id: str,
        db: Session = None
    ) -> bool:
        """
        Delete document.
        
        Features:
        - Document removal
        - File cleanup
        - Vector store cleanup
        """
        try:
            if not db:
                return False
            
            document = db.query(Document).filter(
                Document.id == uuid.UUID(document_id),
                Document.user_id == uuid.UUID(user_id)
            ).first()
            
            if not document:
                raise NotFoundError("Document", document_id)
            
            # Delete physical file
            file_path = os.path.join(self.storage_path, document.filename)
            if os.path.exists(file_path):
                os.remove(file_path)
            
            # Delete from database
            db.delete(document)
            db.commit()
            
            # TODO: Clean up vector embeddings
            
            logger.info(f"Deleted document {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting document: {str(e)}")
            if db:
                db.rollback()
            return False
    
    async def _validate_file(self, file: UploadFile) -> None:
        """Validate uploaded file."""
        # Check file size
        file_content = await file.read()
        if len(file_content) > self.max_file_size:
            raise ValidationError(f"File size exceeds maximum allowed size of {self.max_file_size} bytes")
        
        # Check file type
        if file.content_type not in self.allowed_types:
            raise ValidationError(f"File type {file.content_type} is not supported")
        
        # Reset file pointer
        await file.seek(0)
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename."""
        if not filename:
            return ""
        return os.path.splitext(filename)[1]
    
    async def _extract_from_pdf(self, file_content: bytes) -> ExtractedText:
        """Extract text from PDF file."""
        try:
            import io
            from pypdf import PdfReader
            
            # Create a BytesIO object from the file content
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text = ""
            page_count = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                text += page_text + "\n"
                logger.info(f"Extracted {len(page_text)} chars from page {page_num + 1}")
            
            logger.info(f"Total extracted text length: {len(text)} characters")
            logger.info(f"Text preview (first 200 chars): {text[:200]}...")
            
            metadata = {
                "format": "pdf",
                "pages": page_count,
                "extraction_method": "pypdf"
            }
            
            return ExtractedText(content=text, metadata=metadata, page_count=page_count)
        except Exception as e:
            logger.error(f"Failed to extract PDF content: {str(e)}")
            raise FileProcessingError(f"Failed to extract PDF content: {str(e)}")
    
    async def _extract_from_docx(self, file_content: bytes) -> ExtractedText:
        """Extract text from DOCX file."""
        try:
            import io
            try:
                from docx import Document as DocxDocument
                
                # Create a BytesIO object from the file content
                docx_file = io.BytesIO(file_content)
                doc = DocxDocument(docx_file)
                
                # Extract text from all paragraphs
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                logger.info(f"Total extracted text length: {len(text)} characters")
                logger.info(f"Text preview (first 200 chars): {text[:200]}...")
                
                metadata = {
                    "format": "docx",
                    "extraction_method": "python-docx"
                }
                
                return ExtractedText(content=text, metadata=metadata)
                
            except ImportError:
                logger.warning("python-docx not installed, falling back to basic text extraction")
                # Fallback: try to read as plain text (will likely produce garbage)
                content = file_content.decode('utf-8', errors='ignore')
                metadata = {
                    "format": "docx",
                    "extraction_method": "fallback-text",
                    "warning": "python-docx not available, content may be garbled"
                }
                return ExtractedText(content=content, metadata=metadata)
                
        except Exception as e:
            logger.error(f"Failed to extract DOCX content: {str(e)}")
            raise FileProcessingError(f"Failed to extract DOCX content: {str(e)}")
    
    async def _extract_from_text(self, file_content: bytes) -> ExtractedText:
        """Extract text from plain text file."""
        try:
            content = file_content.decode('utf-8')
            metadata = {
                "format": "text",
                "encoding": "utf-8"
            }
            return ExtractedText(content=content, metadata=metadata)
        except Exception as e:
            raise FileProcessingError(f"Failed to extract text content: {str(e)}")
    
    def _chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ) -> List[str]:
        """Split text into overlapping chunks."""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - chunk_overlap
            
            if start >= len(text):
                break
        
        return chunks
    
    def _document_to_dict(self, document: Document) -> Dict[str, Any]:
        """Convert document model to dictionary."""
        return {
            "id": str(document.id),
            "filename": document.filename,
            "original_filename": document.original_filename,
            "file_hash": document.file_hash,
            "file_size": document.file_size,
            "file_type": document.file_type,
            "user_id": str(document.user_id),
            "created_at": document.created_at.isoformat(),
            "updated_at": document.updated_at.isoformat() if document.updated_at else None
        }
# update Sun Jul  6 02:54:59 IST 2025
